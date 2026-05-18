"""Render GroupXX_milestone_report.md to a Word .docx file via python-docx."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO = Path(__file__).resolve().parent.parent
MD = REPO / "milestone" / "GroupXX_milestone_report.md"
OUT = REPO / "milestone" / "GroupXX_milestone_report.docx"

NAVY = RGBColor(0x0B, 0x3D, 0x91)


def parse_yaml(text):
    if not text.startswith("---"):
        return {}, text
    end = text.find("---", 3)
    yaml_block = text[3:end]
    body = text[end + 3:]
    meta = {}
    for line in yaml_block.splitlines():
        m = re.match(r'(\w+):\s*"?([^"\n]+)"?', line)
        if m:
            meta[m.group(1)] = m.group(2).strip().strip('"')
    return meta, body


def add_runs_inline(para, text):
    """Parse **bold**, *italic*, `code` markers and add styled runs."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*\n]+\*|`[^`]+`)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            para.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith("**"):
            r = para.add_run(token[2:-2]); r.bold = True
        elif token.startswith("`"):
            r = para.add_run(token[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9.5)
        elif token.startswith("*"):
            r = para.add_run(token[1:-1]); r.italic = True
        pos = m.end()
    if pos < len(text):
        para.add_run(text[pos:])


def parse_md_table(lines, i):
    header = [c.strip() for c in lines[i].strip().strip("|").split("|")]
    i += 2  # skip separator
    rows = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
        i += 1
    return header, rows, i


def main():
    raw = MD.read_text(encoding="utf-8")
    meta, body = parse_yaml(raw)

    doc = Document()
    # Set margins to match PDF version
    for section in doc.sections:
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.55)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Title block
    if "title" in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(meta["title"]); r.bold = True; r.font.size = Pt(16)
        r.font.color.rgb = NAVY
    if "subtitle" in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(meta["subtitle"]); r.italic = True; r.font.size = Pt(11)
    if "author" in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(meta["author"]).font.size = Pt(10)
    if "date" in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(meta["date"]).font.size = Pt(10)

    lines = body.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        s = line.strip()

        if s.startswith("# "):
            p = doc.add_paragraph()
            r = p.add_run(s[2:]); r.bold = True; r.font.size = Pt(14)
            r.font.color.rgb = NAVY
            i += 1
        elif s.startswith("## "):
            p = doc.add_paragraph()
            r = p.add_run(s[3:]); r.bold = True; r.font.size = Pt(12)
            r.font.color.rgb = NAVY
            i += 1
        elif s.startswith("### "):
            p = doc.add_paragraph()
            r = p.add_run(s[4:]); r.bold = True; r.font.size = Pt(11)
            i += 1
        elif s.startswith("|") and i + 1 < len(lines) and set(lines[i+1].strip()).issubset(set("|-: ")):
            header, rows, i = parse_md_table(lines, i)
            ncols = len(header)
            table = doc.add_table(rows=1 + len(rows), cols=ncols)
            table.style = "Light Grid Accent 1"
            for c, txt in enumerate(header):
                cell = table.rows[0].cells[c]
                cell.text = ""
                p = cell.paragraphs[0]
                add_runs_inline(p, txt)
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(9)
            for r_idx, row in enumerate(rows, 1):
                for c, txt in enumerate(row[:ncols]):
                    cell = table.rows[r_idx].cells[c]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    add_runs_inline(p, txt)
                    for run in p.runs:
                        run.font.size = Pt(9)
        elif s.startswith("- ") or s.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs_inline(p, s[2:])
            i += 1
        elif re.match(r"^\d+\.\s", s):
            p = doc.add_paragraph(style="List Number")
            add_runs_inline(p, re.sub(r"^\d+\.\s", "", s))
            i += 1
        elif s.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            p = doc.add_paragraph()
            r = p.add_run("\n".join(code_lines))
            r.font.name = "Consolas"; r.font.size = Pt(9)
        elif s == "---":
            doc.add_paragraph()
            i += 1
        elif s == "":
            i += 1
        else:
            # Collect a paragraph (consecutive non-blank lines that aren't structural)
            buf = [s]
            i += 1
            while i < len(lines):
                nxt = lines[i].strip()
                if not nxt or nxt.startswith(("#", "-", "*", "|", "```")) or re.match(r"^\d+\.\s", nxt):
                    break
                buf.append(nxt)
                i += 1
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_runs_inline(p, " ".join(buf))

    try:
        doc.save(str(OUT))
        print("Wrote", OUT)
    except PermissionError:
        tmp = OUT.with_suffix(".new.docx")
        doc.save(str(tmp))
        print(f"WARNING: {OUT.name} locked (open in Word?). Wrote {tmp.name} instead.")


if __name__ == "__main__":
    main()
