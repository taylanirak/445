"""Render Group24_final_report.md to a Word .docx via python-docx.

Mirrors the milestone DOCX builder and substitutes the live ``results.json``
numbers through ``src.report_fill`` (same Markdown subset, no pandoc).
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

FINAL = Path(__file__).resolve().parent
sys.path.insert(0, str(FINAL))
from src.report_fill import load_filled  # noqa: E402

MD = FINAL / "Group24_final_report.md"
OUT = FINAL / "Group24_final_report.docx"
NAVY = RGBColor(0x0B, 0x3D, 0x91)


def parse_yaml(text):
    text = text.lstrip()
    if not text.startswith("---"):
        return {}, text
    end = text.find("---", 3)
    meta = {}
    for line in text[3:end].splitlines():
        m = re.match(r'(\w+):\s*"?([^"\n]+)"?', line)
        if m:
            meta[m.group(1)] = m.group(2).strip().strip('"')
    return meta, text[end + 3:]


def add_runs_inline(para, text):
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*\n]+\*|`[^`]+`)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            para.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith("**"):
            r = para.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("`"):
            r = para.add_run(tok[1:-1]); r.font.name = "Consolas"
            r.font.size = Pt(9.5)
        else:
            r = para.add_run(tok[1:-1]); r.italic = True
        pos = m.end()
    if pos < len(text):
        para.add_run(text[pos:])


def parse_md_table(lines, i):
    header = [c.strip() for c in lines[i].strip().strip("|").split("|")]
    i += 2
    rows = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
        i += 1
    return header, rows, i


def main():
    meta, body = parse_yaml(load_filled(MD))
    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.55)
    base = doc.styles["Normal"]
    base.font.name = "Calibri"
    base.font.size = Pt(10)

    for key, sz, bold, color in [("title", 16, True, NAVY),
                                 ("subtitle", 11, False, None),
                                 ("author", 10, False, None),
                                 ("date", 10, False, None)]:
        if key in meta:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(meta[key]); r.font.size = Pt(sz)
            r.bold = bold
            if key == "subtitle":
                r.italic = True
            if color:
                r.font.color.rgb = color

    lines = body.splitlines()
    i = 0
    while i < len(lines):
        s = lines[i].strip()
        if s.startswith("# "):
            p = doc.add_paragraph(); r = p.add_run(s[2:])
            r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
            i += 1
        elif s.startswith("## "):
            p = doc.add_paragraph(); r = p.add_run(s[3:])
            r.bold = True; r.font.size = Pt(12); r.font.color.rgb = NAVY
            i += 1
        elif s.startswith("### "):
            p = doc.add_paragraph(); r = p.add_run(s[4:])
            r.bold = True; r.font.size = Pt(11)
            i += 1
        elif (s.startswith("|") and i + 1 < len(lines)
              and set(lines[i + 1].strip()).issubset(set("|-: "))):
            header, rows, i = parse_md_table(lines, i)
            ncols = len(header)
            table = doc.add_table(rows=1 + len(rows), cols=ncols)
            table.style = "Light Grid Accent 1"
            for c, txt in enumerate(header):
                cell = table.rows[0].cells[c]
                cell.text = ""
                add_runs_inline(cell.paragraphs[0], txt)
                for run in cell.paragraphs[0].runs:
                    run.bold = True; run.font.size = Pt(9)
            for r_idx, row in enumerate(rows, 1):
                for c, txt in enumerate(row[:ncols]):
                    cell = table.rows[r_idx].cells[c]
                    cell.text = ""
                    add_runs_inline(cell.paragraphs[0], txt)
                    for run in cell.paragraphs[0].runs:
                        run.font.size = Pt(9)
        elif s.startswith("- ") or s.startswith("* "):
            add_runs_inline(doc.add_paragraph(style="List Bullet"), s[2:])
            i += 1
        elif re.match(r"^\d+\.\s", s):
            add_runs_inline(doc.add_paragraph(style="List Number"),
                            re.sub(r"^\d+\.\s", "", s))
            i += 1
        elif s.startswith("```"):
            i += 1
            code = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i]); i += 1
            i += 1
            r = doc.add_paragraph().add_run("\n".join(code))
            r.font.name = "Consolas"; r.font.size = Pt(9)
        elif s == "---":
            doc.add_paragraph(); i += 1
        elif s == "":
            i += 1
        else:
            buf = [s]; i += 1
            while i < len(lines):
                nxt = lines[i].strip()
                if (not nxt or nxt.startswith(("#", "-", "*", "|", "```"))
                        or re.match(r"^\d+\.\s", nxt)):
                    break
                buf.append(nxt); i += 1
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_runs_inline(p, " ".join(buf))

    try:
        doc.save(str(OUT))
        print("Wrote", OUT)
    except PermissionError:
        tmp = OUT.with_suffix(".new.docx")
        doc.save(str(tmp))
        print(f"WARNING: {OUT.name} locked. Wrote {tmp.name} instead.")


if __name__ == "__main__":
    main()
